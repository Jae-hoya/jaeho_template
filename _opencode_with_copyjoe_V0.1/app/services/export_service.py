from io import BytesIO
from html import escape
from typing import Any

from docx import Document as create_document

from app.schemas.copy import CopyGenerateResponse


class ExportService:
    def export_docx(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        document = create_document()
        document.add_heading("Copyjoe Export", level=1)

        self._add_section(document, "Head", result.head)
        self._add_section(document, "Body", result.body)
        self._add_section(document, "CTA", result.cta)
        self._add_section(document, "Slogan", result.slogan)
        self._add_section(document, "SNS", result.sns)
        self._add_section(document, "Description", result.description)
        self._add_section(document, "Rationale", result.rationale)

        document.add_heading("Storyboard", level=2)
        for line in result.storyboard_outline:
            document.add_paragraph(line, style="List Bullet")

        if result.sources:
            document.add_heading("Sources", level=2)
            for source in result.sources:
                base = f"[{source.source_type}] {source.title or ''}".strip()
                if source.url:
                    base += f" ({source.url})"
                document.add_paragraph(base)

        buffer = BytesIO()
        document.save(buffer)

        safe_name = file_name if file_name.endswith(".docx") else f"{file_name}.docx"
        return safe_name, buffer.getvalue()

    def export_markdown(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        lines: list[str] = [
            "# Copyjoe Export",
            "",
            "## Head",
            result.head,
            "",
            "## Body",
            result.body,
            "",
            "## CTA",
            result.cta,
            "",
            "## Slogan",
            result.slogan,
            "",
            "## SNS",
            result.sns,
            "",
            "## Description",
            result.description,
            "",
            "## Rationale",
            result.rationale,
            "",
            "## Storyboard",
        ]

        for item in result.storyboard_outline:
            lines.append(f"- {item}")

        if result.sources:
            lines.append("")
            lines.append("## Sources")
            for source in result.sources:
                title = source.title or "untitled"
                suffix = f" ({source.url})" if source.url else ""
                lines.append(f"- [{source.source_type}] {title}{suffix}")

        content = "\n".join(lines).strip() + "\n"
        safe_name = file_name if file_name.endswith(".md") else f"{file_name}.md"
        return safe_name, content.encode("utf-8")

    def export_doc(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        html_sections = [
            "<html><head><meta charset='utf-8'></head><body>",
            "<h1>Copyjoe Export</h1>",
            f"<h2>Head</h2><p>{escape(result.head)}</p>",
            f"<h2>Body</h2><p>{escape(result.body)}</p>",
            f"<h2>CTA</h2><p>{escape(result.cta)}</p>",
            f"<h2>Slogan</h2><p>{escape(result.slogan)}</p>",
            f"<h2>SNS</h2><p>{escape(result.sns)}</p>",
            f"<h2>Description</h2><p>{escape(result.description)}</p>",
            f"<h2>Rationale</h2><p>{escape(result.rationale)}</p>",
            "<h2>Storyboard</h2><ul>",
        ]

        for line in result.storyboard_outline:
            html_sections.append(f"<li>{escape(line)}</li>")

        html_sections.append("</ul>")

        if result.sources:
            html_sections.append("<h2>Sources</h2><ul>")
            for source in result.sources:
                title = escape(source.title or "untitled")
                if source.url:
                    html_sections.append(f"<li>[{source.source_type}] {title} ({escape(source.url)})</li>")
                else:
                    html_sections.append(f"<li>[{source.source_type}] {title}</li>")
            html_sections.append("</ul>")

        html_sections.append("</body></html>")

        safe_name = file_name if file_name.endswith(".doc") else f"{file_name}.doc"
        return safe_name, "".join(html_sections).encode("utf-8")

    def _add_section(self, document: Any, title: str, text: str) -> None:
        document.add_heading(title, level=2)
        document.add_paragraph(text or "")
