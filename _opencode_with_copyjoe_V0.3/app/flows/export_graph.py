from html import escape
from io import BytesIO
from typing import Any, Literal, TypedDict

from docx import Document as create_document
from langgraph.graph import END, START, StateGraph

from app.schemas.copy import CopyGenerateResponse


class ExportState(TypedDict):
    action: Literal["docx", "md", "doc"]
    file_name: str
    result: CopyGenerateResponse
    safe_name: str
    content: bytes
    output: tuple[str, bytes] | None


class ExportWorkflowGraph:
    def __init__(self) -> None:
        builder = StateGraph(ExportState)
        builder.add_node("prepare", self._prepare_name)
        builder.add_node("build_docx", self._build_docx)
        builder.add_node("build_markdown", self._build_markdown)
        builder.add_node("build_doc", self._build_doc)
        builder.add_node("finalize", self._finalize_output)

        builder.add_edge(START, "prepare")
        builder.add_conditional_edges(
            "prepare",
            self._route_after_prepare,
            {
                "build_docx": "build_docx",
                "build_markdown": "build_markdown",
                "build_doc": "build_doc",
            },
        )
        builder.add_edge("build_docx", "finalize")
        builder.add_edge("build_markdown", "finalize")
        builder.add_edge("build_doc", "finalize")
        builder.add_edge("finalize", END)
        self._graph = builder.compile()

    def run_docx(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        return self._run(action="docx", file_name=file_name, result=result)

    def run_markdown(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        return self._run(action="md", file_name=file_name, result=result)

    def run_doc(self, file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        return self._run(action="doc", file_name=file_name, result=result)

    def _run(self, action: Literal["docx", "md", "doc"], file_name: str, result: CopyGenerateResponse) -> tuple[str, bytes]:
        state = self._graph.invoke(
            {
                "action": action,
                "file_name": file_name,
                "result": result,
                "safe_name": "",
                "content": b"",
                "output": None,
            }
        )
        output = state.get("output")
        if output is None:
            raise RuntimeError("Export graph failed to build output")
        return output

    def _prepare_name(self, state: ExportState) -> dict[str, object]:
        suffix_by_action = {
            "docx": ".docx",
            "md": ".md",
            "doc": ".doc",
        }
        suffix = suffix_by_action[state["action"]]
        file_name = state["file_name"]
        safe_name = file_name if file_name.endswith(suffix) else f"{file_name}{suffix}"
        return {"safe_name": safe_name}

    def _route_after_prepare(self, state: ExportState) -> str:
        if state["action"] == "docx":
            return "build_docx"
        if state["action"] == "md":
            return "build_markdown"
        return "build_doc"

    def _build_docx(self, state: ExportState) -> dict[str, object]:
        result = state["result"]
        document = create_document()
        document.add_heading("Copyjoe Export", level=1)

        self._add_section(document, "Head", result.head)
        self._add_section(document, "Body", result.body)
        self._add_section(document, "CTA", result.cta)
        self._add_section(document, "Slogan", result.slogan)
        self._add_section(document, "SNS", result.sns)
        self._add_section(document, "Description", result.description)
        self._add_section(document, "Rationale", result.rationale)

        document.add_heading("Storyboard", level=2)
        for line in result.storyboard_outline:
            document.add_paragraph(line, style="List Bullet")

        if result.sources:
            document.add_heading("Sources", level=2)
            for source in result.sources:
                base = f"[{source.source_type}] {source.title or ''}".strip()
                if source.url:
                    base += f" ({source.url})"
                document.add_paragraph(base)

        buffer = BytesIO()
        document.save(buffer)
        return {"content": buffer.getvalue()}

    def _build_markdown(self, state: ExportState) -> dict[str, object]:
        result = state["result"]
        lines: list[str] = [
            "# Copyjoe Export",
            "",
            "## Head",
            result.head,
            "",
            "## Body",
            result.body,
            "",
            "## CTA",
            result.cta,
            "",
            "## Slogan",
            result.slogan,
            "",
            "## SNS",
            result.sns,
            "",
            "## Description",
            result.description,
            "",
            "## Rationale",
            result.rationale,
            "",
            "## Storyboard",
        ]

        for item in result.storyboard_outline:
            lines.append(f"- {item}")

        if result.sources:
            lines.append("")
            lines.append("## Sources")
            for source in result.sources:
                title = source.title or "untitled"
                suffix = f" ({source.url})" if source.url else ""
                lines.append(f"- [{source.source_type}] {title}{suffix}")

        content = "\n".join(lines).strip() + "\n"
        return {"content": content.encode("utf-8")}

    def _build_doc(self, state: ExportState) -> dict[str, object]:
        result = state["result"]
        html_sections = [
            "<html><head><meta charset='utf-8'></head><body>",
            "<h1>Copyjoe Export</h1>",
            f"<h2>Head</h2><p>{escape(result.head)}</p>",
            f"<h2>Body</h2><p>{escape(result.body)}</p>",
            f"<h2>CTA</h2><p>{escape(result.cta)}</p>",
            f"<h2>Slogan</h2><p>{escape(result.slogan)}</p>",
            f"<h2>SNS</h2><p>{escape(result.sns)}</p>",
            f"<h2>Description</h2><p>{escape(result.description)}</p>",
            f"<h2>Rationale</h2><p>{escape(result.rationale)}</p>",
            "<h2>Storyboard</h2><ul>",
        ]

        for line in result.storyboard_outline:
            html_sections.append(f"<li>{escape(line)}</li>")

        html_sections.append("</ul>")

        if result.sources:
            html_sections.append("<h2>Sources</h2><ul>")
            for source in result.sources:
                title = escape(source.title or "untitled")
                if source.url:
                    html_sections.append(f"<li>[{source.source_type}] {title} ({escape(source.url)})</li>")
                else:
                    html_sections.append(f"<li>[{source.source_type}] {title}</li>")
            html_sections.append("</ul>")

        html_sections.append("</body></html>")
        return {"content": "".join(html_sections).encode("utf-8")}

    def _finalize_output(self, state: ExportState) -> dict[str, object]:
        return {
            "output": (state["safe_name"], state["content"]),
        }

    def _add_section(self, document: Any, title: str, text: str) -> None:
        document.add_heading(title, level=2)
        document.add_paragraph(text or "")
