import argparse
import json
import os
import subprocess
import sys
import time
from pathlib import Path
from typing import Any

import requests
from docx import Document


def _start_server(port: int, extra_env: dict[str, str]) -> subprocess.Popen[Any]:
    env = os.environ.copy()
    env.update(extra_env)
    return subprocess.Popen(
        [sys.executable, "-m", "uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", str(port)],
        env=env,
    )


def _wait_until_ready(base_url: str, timeout_sec: int = 45) -> None:
    started = time.time()
    while time.time() - started < timeout_sec:
        try:
            response = requests.get(f"{base_url}/health", timeout=5)
            if response.status_code == 200:
                return
        except Exception:
            pass
        time.sleep(0.7)
    raise RuntimeError("FastAPI server did not become ready in time")


def _post_json(base_url: str, path: str, payload: dict[str, Any], timeout_sec: int = 90) -> requests.Response:
    return requests.post(f"{base_url}{path}", json=payload, timeout=timeout_sec)


def _is_mock_like(copy_payload: dict[str, Any], generated: dict[str, Any]) -> bool:
    expected_prefix = f"{copy_payload['target_audience']}를 위한 {copy_payload['product_name']}"
    return str(generated.get("head", "")).startswith(expected_prefix)


def _ollama_available(base_url: str) -> bool:
    try:
        response = requests.get(f"{base_url.rstrip('/')}/api/tags", timeout=8)
        return response.status_code == 200
    except Exception:
        return False


def run_checks(provider: str, port: int, ollama_model: str) -> dict[str, Any]:
    provider = provider.strip().lower()
    if provider not in {"openai", "ollama"}:
        raise ValueError("provider must be openai or ollama")

    extra_env: dict[str, str] = {
        "LLM_PROVIDER": provider,
        "FORCE_MOCK_MODE": "false",
    }

    if provider == "ollama":
        extra_env["OLLAMA_CHAT_MODEL"] = ollama_model

    process = _start_server(port=port, extra_env=extra_env)
    base_url = f"http://127.0.0.1:{port}"

    try:
        _wait_until_ready(base_url)

        summary: dict[str, Any] = {
            "provider": provider,
            "base_url": base_url,
            "checks": {},
        }

        for path in ["/", "/health", "/openapi.json", "/docs", "/api/v1/meta/copy-form-guide", "/api/v1/history/threads"]:
            response = requests.get(f"{base_url}{path}", timeout=20)
            summary["checks"][path] = response.status_code

        health = requests.get(f"{base_url}/health", timeout=20).json()
        summary["health"] = health

        copy_payload = {
            "product_name": "Copyjoe",
            "target_audience": "퍼포먼스 마케터",
            "pain_point": "카피 작성에 시간이 오래 걸린다",
            "differentiator": "RAG와 Tavily를 결합한 근거 기반 생성",
            "tone": "신뢰형",
            "objective": "click",
            "styles": ["head", "body", "cta", "slogan", "sns", "description"],
            "channel": "상세페이지",
            "language": "ko",
            "web_search_mode": False,
            "use_rag": False,
            "top_k": 5,
        }

        copy_response = _post_json(base_url, "/api/v1/copy/generate", copy_payload, timeout_sec=120)
        summary["checks"]["copy_generate"] = copy_response.status_code
        copy_body = copy_response.json()
        summary["copy_non_empty"] = {
            key: bool((copy_body.get(key) or "").strip())
            for key in ["head", "body", "cta", "slogan", "sns", "description"]
        }
        summary["copy_mock_detected"] = _is_mock_like(copy_payload, copy_body)

        history_create = _post_json(base_url, "/api/v1/history/threads", {"title": "full-check-thread"}, timeout_sec=60)
        summary["checks"]["history_create"] = history_create.status_code
        thread_id = history_create.json().get("thread_id")

        history_message = _post_json(
            base_url,
            f"/api/v1/history/threads/{thread_id}/messages",
            {"role": "user", "content": "테스트 메시지", "metadata": {"source": "full-check"}},
            timeout_sec=60,
        )
        summary["checks"]["history_append_message"] = history_message.status_code

        history_detail = requests.get(f"{base_url}/api/v1/history/threads/{thread_id}", timeout=60)
        summary["checks"]["history_get_thread"] = history_detail.status_code

        copy_prompt_response = _post_json(
            base_url,
            "/api/v1/copy/generate",
            {
                "prompt": "클릭률이 떨어져서 근거 중심 카피를 빠르게 만들고 싶다",
                "styles": ["head", "cta", "sns"],
                "language": "english",
            },
            timeout_sec=120,
        )
        summary["checks"]["copy_generate_prompt_mode"] = copy_prompt_response.status_code
        copy_prompt_body = copy_prompt_response.json()
        summary["copy_prompt_language"] = copy_prompt_body.get("normalized_request", {}).get("language")

        sample = Path("tmp_full_check_sample.docx")
        document = Document()
        document.add_heading("Copyjoe RAG Input", level=1)
        document.add_paragraph("Copyjoe는 업로드 문서를 변환하고 RAG 검색으로 생성 근거를 강화한다.")
        document.add_paragraph("Tavily 웹 검색은 최신 시장 문맥을 보완한다.")
        document.save(str(sample))

        with sample.open("rb") as file_obj:
            upload_response = requests.post(
                f"{base_url}/api/v1/files/upload",
                files=[
                    (
                        "files",
                        (
                            sample.name,
                            file_obj,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
                timeout=60,
            )

        summary["checks"]["upload"] = upload_response.status_code
        upload_body = upload_response.json()
        doc_ids = [item["document_id"] for item in upload_body.get("files", []) if item.get("success")]
        summary["uploaded_doc_ids"] = len(doc_ids)

        index_response = _post_json(
            base_url,
            "/api/v1/rag/index",
            {"document_ids": doc_ids, "chunk_size": 500, "chunk_overlap": 80},
        )
        summary["checks"]["rag_index"] = index_response.status_code
        summary["indexed"] = index_response.json()

        search_response = _post_json(
            base_url,
            "/api/v1/rag/search",
            {"query": "Copyjoe에서 Tavily의 역할은?", "top_k": 3},
        )
        summary["checks"]["rag_search"] = search_response.status_code
        search_body = search_response.json()
        summary["rag_results"] = len(search_body.get("results", []))

        copy_payload["use_rag"] = True
        copy_rag_response = _post_json(base_url, "/api/v1/copy/generate", copy_payload, timeout_sec=120)
        summary["checks"]["copy_generate_rag"] = copy_rag_response.status_code
        copy_rag_body = copy_rag_response.json()
        summary["copy_rag_sources"] = len(copy_rag_body.get("sources", []))

        web_search_response = _post_json(
            base_url,
            "/api/v1/web/search",
            {"query": "copywriting saas landing page", "max_results": 3},
            timeout_sec=120,
        )
        summary["checks"]["web_search"] = web_search_response.status_code
        web_search_body = web_search_response.json()
        summary["web_results"] = len(web_search_body.get("results", []))

        landing_query_response = _post_json(
            base_url,
            "/api/v1/web/landing/analyze",
            {"query": "copywriting saas landing page", "max_results": 3},
            timeout_sec=150,
        )
        summary["checks"]["landing_analyze_query"] = landing_query_response.status_code
        landing_query_body = landing_query_response.json()
        summary["landing_query"] = {
            "url": landing_query_body.get("url"),
            "from_tavily": landing_query_body.get("from_tavily"),
            "h1": len(landing_query_body.get("h1", [])),
            "h2": len(landing_query_body.get("h2", [])),
            "cta": len(landing_query_body.get("cta_buttons", [])),
            "body_len": len(landing_query_body.get("body", "")),
        }

        landing_url_response = _post_json(
            base_url,
            "/api/v1/web/landing/analyze",
            {"url": "https://example.com"},
            timeout_sec=120,
        )
        summary["checks"]["landing_analyze_url"] = landing_url_response.status_code
        landing_url_body = landing_url_response.json()
        summary["landing_url"] = {
            "url": landing_url_body.get("url"),
            "from_tavily": landing_url_body.get("from_tavily"),
            "h1": len(landing_url_body.get("h1", [])),
            "h2": len(landing_url_body.get("h2", [])),
            "cta": len(landing_url_body.get("cta_buttons", [])),
            "body_len": len(landing_url_body.get("body", "")),
        }

        export_response = _post_json(
            base_url,
            "/api/v1/export/docx",
            {"file_name": "copyjoe_check.docx", "result": copy_body},
            timeout_sec=60,
        )
        summary["checks"]["export_docx"] = export_response.status_code
        summary["export_docx_len"] = len(export_response.content)

        export_md_response = _post_json(
            base_url,
            "/api/v1/export/md",
            {"file_name": "copyjoe_check.md", "result": copy_body},
            timeout_sec=60,
        )
        summary["checks"]["export_md"] = export_md_response.status_code
        summary["export_md_len"] = len(export_md_response.content)

        export_doc_response = _post_json(
            base_url,
            "/api/v1/export/doc",
            {"file_name": "copyjoe_check.doc", "result": copy_body},
            timeout_sec=60,
        )
        summary["checks"]["export_doc"] = export_doc_response.status_code
        summary["export_doc_len"] = len(export_doc_response.content)

        summary["ollama_http_available"] = _ollama_available(os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"))

        try:
            sample.unlink(missing_ok=True)
        except Exception:
            pass

        return summary
    finally:
        process.terminate()
        try:
            process.wait(timeout=15)
        except Exception:
            process.kill()


def main() -> int:
    parser = argparse.ArgumentParser(description="Run full Copyjoe API checks")
    parser.add_argument("--provider", choices=["openai", "ollama"], default="openai")
    parser.add_argument("--port", type=int, default=8012)
    parser.add_argument("--ollama-model", default="qwen3:8b")
    args = parser.parse_args()

    result = run_checks(provider=args.provider, port=args.port, ollama_model=args.ollama_model)
    print(json.dumps(result, ensure_ascii=False, indent=2))

    required = [
        "/",
        "/health",
        "/openapi.json",
        "/docs",
        "/api/v1/meta/copy-form-guide",
        "/api/v1/history/threads",
        "copy_generate",
        "copy_generate_prompt_mode",
        "history_create",
        "history_append_message",
        "history_get_thread",
        "upload",
        "rag_index",
        "rag_search",
        "copy_generate_rag",
        "web_search",
        "landing_analyze_query",
        "landing_analyze_url",
        "export_docx",
        "export_md",
        "export_doc",
    ]

    failed = [name for name in required if result["checks"].get(name) != 200]
    if failed:
        print(json.dumps({"status": "failed", "failed_checks": failed}, ensure_ascii=False))
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
